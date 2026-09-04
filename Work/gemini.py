# Save this script as generate_docs.py and run: python generate_docs.py
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def create_blueprint_document():
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("AAPL Sales & Operations Portal")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(13, 110, 253)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run("Technical Architecture & System Design Blueprint (Modular Roadmap)")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(108, 117, 125)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    def add_section_header(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(33, 37, 41)

    # 1. Executive Summary
    add_section_header("1. Executive Summary & Core Objectives")
    doc.add_paragraph(
        "The AAPL Portal is transitioning from a single-file Streamlit application into a decoupled, "
        "modular enterprise platform. The new architecture integrates with daily restored Shoper SQL Databases, "
        "automates live Tally ERP 9 / Prime synchronizations, enforces role-based access control, tracks dispatch SLAs, "
        "and automates debtor collections."
    )

    # 2. Directory Structure
    add_section_header("2. Modular Directory Architecture")
    p_code = doc.add_paragraph()
    r_code = p_code.add_run(
"""aapl_portal/
├── config/ (settings.py)
├── database/ (connection.py, models.py)
├── services/ (tally_service.py, shoper_service.py, delivery_service.py)
├── schedulers/ (reminder_scheduler.py)
├── ui/ (auth.py, components/, views/)
├── app.py (Streamlit Router)
└── sync_runner.py (ETL Runner)"""
    )
    r_code.font.name = 'Consolas'
    r_code.font.size = Pt(9.5)

    # 3. RBAC Table
    add_section_header("3. Role-Based Access Control (RBAC) Matrix")
    headers = ["Portal View", "Admin", "Manager", "Salesman", "Dispatch"]
    data = [
        ["Executive Dashboard", "Full Access", "Full Access", "No Access", "No Access"],
        ["Salesman 360° View", "All Outlets", "All Outlets", "Assigned", "No Access"],
        ["Delivery Dashboard", "View / Edit", "View / Edit", "View Only", "View / Edit"],
        ["Outstanding Debtors", "Full Access", "Full Access", "Assigned", "No Access"],
        ["Stock Details", "View / Export", "View / Export", "View / Export", "View / Export"],
        ["Investment Master", "Full Access", "No Access", "No Access", "No Access"]
    ]

    table = doc.add_table(rows=len(data) + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], "0D6EFD")

    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F8F9FA")

    # 4. Delivery SLA Matrix
    add_section_header("4. Delivery SLA Aging Matrix")
    sla_data = [
        ["0 Days (Same Day)", "Neutral Gray", "On schedule (Standard operational turnaround)"],
        ["1 - 2 Days", "Warning Yellow", "Approaching dispatch SLA limit"],
        ["> 2 Days", "Alert Red", "Critical dispatch bottleneck requiring escalation"]
    ]
    sla_table = doc.add_table(rows=4, cols=3)
    sla_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, h in enumerate(["Aging Threshold", "Indicator", "Operational Meaning"]):
        sla_table.rows[0].cells[i].text = h
        sla_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        sla_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(sla_table.rows[0].cells[i], "495057")

    for row_idx, row_data in enumerate(sla_data):
        r_cells = sla_table.rows[row_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            r_cells[c_idx].text = val

    # Save
    doc.save("AAPL_Sales_Portal_Architecture_Blueprint.docx")
    print("Document successfully created: AAPL_Sales_Portal_Architecture_Blueprint.docx")

if __name__ == "__main__":
    create_blueprint_document()